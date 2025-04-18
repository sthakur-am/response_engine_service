import os
import io
import base64

import requests

from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient
from langchain_community.document_loaders import AzureBlobStorageContainerLoader

from docx import Document
from pypdf import PdfReader

class DocumentBase(object):
    def __init__(self, config):
        self.access_key = config["STORAGE_ACCESS_KEY"]
        self.connection_str = config["STORAGE_CONNECTION_STR"]
        self.container_name = 'rfpengine'

    def upload_file(self, filename, content):
        blob_service_client = BlobServiceClient.from_connection_string(self.connection_str)

        # Create ContainerClient
        container_client = blob_service_client.get_container_client(self.container_name)

        # Create BlobClient
        blob_client = container_client.get_blob_client(filename)

        # Upload the file to the blob
        blob_client.upload_blob(content, overwrite=True)

    def extract_text_from_blob(self, filename):
        blob_service_client = BlobServiceClient.from_connection_string(self.connection_str)

        # Create ContainerClient
        container_client = blob_service_client.get_container_client(self.container_name)

        # Create BlobClient
        blob_client = container_client.get_blob_client(filename)

        stream = io.BytesIO()
        num_bytes = blob_client.download_blob().readinto(stream)
        print(f"Number of bytes: {num_bytes}")

        filetype = filename.split(".")[1]
        if filetype in ["pdf"]:
            return PDFReader(stream).read_text()
        elif filetype in ["doc", "docx"]:
            return WordReader(stream).read_text()
        else:
            raise Exception("Unknown file type ", filetype);

    def read_blob_as_base64(self, filename):
        blob_service_client = BlobServiceClient.from_connection_string(self.connection_str)

        # Create ContainerClient
        container_client = blob_service_client.get_container_client(self.container_name)

        # Create BlobClient
        blob_client = container_client.get_blob_client(filename)

        # stream = io.BytesIO()
        blob_data = blob_client.download_blob().readall()
        # print(f"Number of bytes: {num_bytes}")

        filetype = filename.split(".")[1]
        if filetype in ["dotx"]:
            return base64.b64encode(blob_data)
        else:
            raise Exception("Unknown file type ", filetype);

    def extract_text_from_document(self, filename):
        full_text = []
        doc = self.load_document(filename)
        for para in doc.paragraphs:
            full_text.append(para.text)

        return '\n'.join(full_text)
        
    def load_document(self, filename):
        doc_filepath = filename
        # self.download_filename(filename)
        if doc_filepath:
            document = Document(filename)
            return document
        else:
            return None
        
    def download_filename(self, filename):
        file_url = f'https://graph.microsoft.com/v1.0/me/drive/root:/{filename}'
        download_filepath = os.path.join(self.working_directory, filename)
        # Make a GET request to download the file
        response = requests.get(file_url, headers=self.headers)

        if response.status_code == 200:
            with open(download_filepath, 'wb') as file:
                file.write(response.content)
            print('File downloaded successfully')
            return download_filepath
        else:
            print(f'Error: {response.status_code}')
            return None
    
    def get_item_id(self, filename):
        items = self.client.item(drive="me", id="root").children.get()

        for item in items:
            if item.name == filename:
                return item.id
            
        return None
    
    def delete_blob_path(self, path):
        blob_service_client = BlobServiceClient.from_connection_string(self.connection_str)
        container_client = blob_service_client.get_container_client(self.container_name)
        blob_list = container_client.list_blobs(path)
        blob_names = [blob.name for blob in blob_list]
        # print(blob_names)
        container_client.delete_blobs(*blob_names)

    def get_users(self, username):
        users_url = "https://graph.microsoft.com/v1.0/me/drive"
        users_response = requests.get(users_url, headers=self.headers)
        users_response.raise_for_status()
        users = users_response.json().get('value', [])
        
        # Find the specific user
        user_id = None
        for user in users:
            print(user)
            # if user['mail'] == username:
            #     user_id = user['id']
            #     break
        
        # if not user_id:
        #     raise Exception("User not found")

    def load_documents(self, path):
        loader = AzureBlobStorageContainerLoader(conn_str=self.connection_str, container=self.container_name, prefix=path)
        documents = loader.load()

        return documents

class PDFReader(object):
    def __init__(self, stream):
        self.reader = PdfReader(stream)
        
    def read_text(self):
        full_text = []
        for page_num in range(self.reader.get_num_pages()):
            page = self.reader.get_page(page_num)
            page_content = page.extract_text()
            full_text.append(page_content)

        return "\n".join(full_text)
    
class WordReader(object):

    def __init__(self, stream):
        self.doc = Document(stream)
        
    def read_text(self):
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)

        return "\n".join(full_text)